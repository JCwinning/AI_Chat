import streamlit as st
import re
import openai
import sys

# import anthropic
import requests
from PIL import Image
import io
from io import BytesIO
import base64
from datetime import datetime
import json
import config
import threading
import queue
import time
import os
from markitdown import MarkItDown
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass  # python-dotenv not installed
import dashscope
from dashscope import MultiModalConversation
import mimetypes
from openpyxl import load_workbook
import csv
import pandas as pd
from docx import Document
from search_providers import get_search_manager, SearchResponse, SearchResult

# File persistence constants
CHAT_HISTORY_FILE = "chat_history.json"
SETTINGS_FILE = "settings.json"
ENV_FILE = ".env"


def load_chat_history():
    """Load chat history from file"""
    try:
        if os.path.exists(CHAT_HISTORY_FILE):
            with open(CHAT_HISTORY_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception as e:
        print(f"Error loading chat history: {e}")
    return {}


def save_chat_history_to_file():
    """Save chat history to file"""
    try:
        # Convert all chat histories to JSON-serializable format before saving
        serializable_chat_sessions = {}
        for chat_id, chat_data in st.session_state.chat_sessions.items():
            chat_copy = chat_data.copy()
            chat_copy["history"] = prepare_chat_history_for_saving(chat_data["history"])
            serializable_chat_sessions[chat_id] = chat_copy

        with open(CHAT_HISTORY_FILE, "w", encoding="utf-8") as f:
            json.dump(serializable_chat_sessions, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Error saving chat history: {e}")


def load_api_keys():
    """Load API keys from .env file or environment"""
    api_keys = {
        "openrouter": "",
        "modelscope": "",
        "siliconflow": "",
        "dashscope": "",
        "bigmodel": "",
        "tavily_api_key": ""
    }

    # Try to load from .env file
    if os.path.exists(ENV_FILE):
        try:
            with open(ENV_FILE, "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if line and "=" in line and not line.startswith("#"):
                        key, value = line.split("=", 1)
                        api_keys[key.strip()] = value.strip()
        except Exception as e:
            print(f"Error loading .env file: {e}")

    # Override with environment variables if they exist
    for key in api_keys:
        env_value = os.getenv(key)
        if env_value:
            api_keys[key] = env_value

    return api_keys


def save_api_keys_to_env(keys_dict):
    """Save API keys to .env file"""
    try:
        # Read existing .env file
        existing_lines = []
        if os.path.exists(ENV_FILE):
            with open(ENV_FILE, "r", encoding="utf-8") as f:
                existing_lines = f.readlines()

        # Parse existing keys and preserve comments
        env_data = {}
        for line in existing_lines:
            line_stripped = line.strip()
            if line_stripped.startswith("#") or not line_stripped:
                continue
            if "=" in line_stripped:
                key, value = line_stripped.split("=", 1)
                env_data[key.strip()] = value.strip()

        # Update with new keys
        env_data.update(keys_dict)

        # Write back to .env file
        with open(ENV_FILE, "w", encoding="utf-8") as f:
            f.write("# .env file\n")
            for key, value in env_data.items():
                if value:  # Only save non-empty keys
                    f.write(f"{key}={value}\n\n")

        # Reload environment
        if 'load_dotenv' in globals():
            load_dotenv(override=True)
        return True
    except Exception as e:
        print(f"Error saving API keys: {e}")
        return False


def load_settings():
    """Load user settings from file"""
    default_settings = {
        "current_models": [],  # Start with empty list - will be populated when keys are added
        "temperature": 1.0,
        "system_prompt": "You are a helpful assistant.",
        "language": "EN",
        "web_search_enabled": "off"  # Options: "off", "on"
    }
    try:
        if os.path.exists(SETTINGS_FILE):
            with open(SETTINGS_FILE, "r", encoding="utf-8") as f:
                saved_settings = json.load(f)
                # Update defaults with saved settings (preserves new keys in defaults)
                default_settings.update(saved_settings)
                
                # Filter out any models that are no longer available
                available_models_set = set(get_all_models().keys())
                if "current_models" in default_settings:
                    filtered_models = [m for m in default_settings["current_models"]
                                     if m in available_models_set]
                    # If all models were filtered out, use empty list
                    default_settings["current_models"] = filtered_models
    except Exception as e:
        print(f"Error loading settings: {e}")
    return default_settings


def save_settings():
    """Save user settings to file"""
    try:
        # Get values from widget states if available, otherwise use session state
        current_models = st.session_state.get("model_selector", st.session_state.current_models)
        temperature = st.session_state.get("temperature_widget", st.session_state.temperature)
        #system_prompt = st.session_state.get("system_prompt_widget", st.session_state.system_prompt)

        settings = {
            "current_models": current_models,
            "temperature": temperature,
            #"system_prompt": system_prompt,
            "language": st.session_state.language,
            "web_search_enabled": st.session_state.get("web_search_enabled", "auto")
        }
        with open(SETTINGS_FILE, "w", encoding="utf-8") as f:
            json.dump(settings, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Error saving settings: {e}")



# Language translations
translations = {
    "EN": {
        "app_title": "🤖 AI Chat",
        "config": "⚙️ Configuration",
        "select_models": "Select Models (choose one or more)",
        "parameters": "Parameters",
        "temperature": "Temperature",
        "temperature_help": "Controls randomness in responses. Lower values (0.0-0.5) = more focused, deterministic responses. Higher values (1.5-1.5) = more creative, varied, and unpredictable responses.",
        "system_prompt": "System Prompt",
        "system_prompt_help": "Instructions for the AI's behavior",
        "upload_images": "📎 Upload Files",
        "choose_images": "Choose files",
        "upload_help": "Upload files (images, PDFs, Word, CSV, TXT, etc.) to include in your message",
        "uploaded_images_title": "**Uploaded Files:**",
        "chat_history": "💬 Chat History",
        "saved_chats": "**Saved Chats:**",
        "no_saved_chats": "No saved chats yet",
        "chat_interface": "💬 Chat Interface",
        "new_chat": "➕ New Chat",
        "delete_chat": "Delete chat",
        "type_message": "Type your message here...",
        "please_select_model": "❌ Please select at least one model.",
        "web_search": "Web Search",
        "web_search_help": "Enable web search to get current information from internet",
        "search_mode_off": "Off",
        "search_mode_always": "On",
        "searching": "🔍 Searching...",
        "search_results": "Search Results",
        "no_search_results": "No search results found",
        "search_error": "Search error occurred",
        "language": "Language",
        "keys": "🔑 Keys",
        "api_keys": "API Keys Management",
        "api_keys_help": "Configure API keys for various models and services",
        "openrouter_key": "OpenRouter API Key",
        "modelscope_key": "ModelScope API Key",
        "siliconflow_key": "SiliconFlow API Key",
        "dashscope_key": "DashScope API Key",
        "bigmodel_key": "BigModel API Key",
        "tavily_key": "Tavily Search API Key",
        "save_keys": "Save Keys",
        "keys_saved": "API keys saved successfully!",
        "keys_env_file": "Keys are stored in .env file",
        "keys_keyring": "Keys are stored securely in system keyring"
    },
    "中文": {
        "app_title": "🤖 AI万答",
        "config": "⚙️ 配置",
        "select_models": "选择模型（可选择多个）",
        "parameters": "参数",
        "temperature": "温度",
        "temperature_help": "控制响应的随机性。较低值（0.0-0.5）= 更专注、确定的响应。较高值（1.0-1.5）= 更有创意、多变和不可预测的响应。",
        "system_prompt": "系统提示",
        "system_prompt_help": "AI 行为指令",
        "upload_images": "📎 上传文件",
        "choose_images": "选择文件",
        "upload_help": "上传文件（图片、PDF、Word、CSV、TXT等）以包含在您的消息中",
        "uploaded_images_title": "**已上传文件：**",
        "chat_history": "💬 聊天历史",
        "saved_chats": "**保存的聊天：**",
        "no_saved_chats": "暂无保存的聊天",
        "chat_interface": "💬 聊天界面",
        "new_chat": "➕ 新聊天",
        "delete_chat": "删除聊天",
        "type_message": "在此输入您的消息...",
        "please_select_model": "❌ 请至少选择一个模型。",
        "web_search": "网络搜索",
        "web_search_help": "启用网络搜索以获取最新信息",
        "search_mode_off": "关闭",
        "search_mode_always": "开启",
        "searching": "🔍 搜索中...",
        "search_results": "搜索结果",
        "no_search_results": "未找到搜索结果",
        "search_error": "搜索出错",
        "language": "语言",
        "keys": "🔑 密钥",
        "api_keys": "API密钥管理",
        "api_keys_help": "配置各种模型和服务的API密钥",
        "openrouter_key": "OpenRouter API密钥",
        "modelscope_key": "ModelScope API密钥",
        "siliconflow_key": "SiliconFlow API密钥",
        "dashscope_key": "DashScope API密钥",
        "bigmodel_key": "BigModel API密钥",
        "tavily_key": "Tavily搜索API密钥",
        "save_keys": "保存密钥",
        "keys_saved": "API密钥保存成功！",
        "keys_env_file": "密钥存储在.env文件中",
        "keys_keyring": "密钥安全存储在系统密钥环中"
    }
}


def t(key):
    """Get translated text for current language"""
    return translations[st.session_state.language][key]

def get_all_models():
    """Get all available models from config.py dynamically plus custom models that have API keys configured"""
    models = {}

    # Load API keys from environment
    api_keys = load_api_keys()

    # Get all attributes from config module except internal ones
    model_attributes = [attr for attr in dir(config) if not attr.startswith("_")]

    for attr in model_attributes:
        if hasattr(config, attr):
            model_config = getattr(config, attr)

            # Check if it's a valid model config (dict with required keys)
            if (
                isinstance(model_config, dict)
                and "base_url" in model_config
                and "models" in model_config
            ):
                # Create a copy to avoid modifying the original
                model_config = model_config.copy()
                # Try to get API key from environment or saved keys
                service_name = attr.lower()

                # Map service names to API key keys
                key_mapping = {
                    "openrouter": "openrouter",
                    "modelscope": "modelscope",
                    "siliconflow": "siliconflow",
                    "dashscope": "dashscope",
                    "bigmodel": "bigmodel"
                }

                # Check if this model has a corresponding API key
                has_api_key = False

                # First check if the model already has an API key assigned in config
                if model_config.get("api_key") and model_config["api_key"].strip():
                    # Check if this assigned API key is available in our saved keys
                    for key_name, key_value in api_keys.items():
                        if key_value and key_value.strip() and key_value == model_config["api_key"]:
                            has_api_key = True
                            break

                # If no API key assigned in config, try to match by service name
                if not has_api_key:
                    for config_name, key_name in key_mapping.items():
                        if config_name in attr.lower():
                            if api_keys.get(key_name):
                                model_config["api_key"] = api_keys[key_name]
                                has_api_key = True
                            break

                # Only add model if it has an API key configured
                if has_api_key:
                    # Use attribute name directly as display name
                    display_name = attr.replace("_", " ").title()
                    models[display_name] = model_config

    # Add custom models from session state (only if they have API keys)
    if "custom_models" in st.session_state:
        for model_name, model_config in st.session_state.custom_models.items():
            # Only add custom models that have API keys
            if model_config.get("api_key"):
                models[model_name] = model_config

    return models


def add_custom_model(model_name, api_key, base_url, models):
    """Add a custom model to session state"""
    if "custom_models" not in st.session_state:
        st.session_state.custom_models = {}

    # Parse models list (comma-separated)
    model_list = [model.strip() for model in models.split(",")]

    # Use the first model as the primary model for API calls
    primary_model = model_list[0] if model_list else models

    st.session_state.custom_models[model_name] = {
        "api_key": api_key,
        "base_url": base_url,
        "models": primary_model,
        "model_list": model_list  # Store full list for potential future use
    }


def create_client(model_config):
    """Create appropriate client based on model configuration"""
    if not model_config.get("api_key"):
        return None

    if "modelscope.cn" in model_config["base_url"]:
        return openai.OpenAI(
            api_key=model_config["api_key"], base_url=model_config["base_url"]
        )
    elif "openrouter.ai" in model_config["base_url"]:
        return openai.OpenAI(
            api_key=model_config["api_key"], base_url=model_config["base_url"]
        )
    else:
        return openai.OpenAI(
            api_key=model_config["api_key"], base_url=model_config["base_url"]
        )


def encode_image_to_base64(image):
    """Encode PIL Image to base64 string"""
    buffered = io.BytesIO()
    image.save(buffered, format="PNG")
    return base64.b64encode(buffered.getvalue()).decode("utf-8")


def decode_base64_to_image(base64_string):
    """Convert base64 string back to PIL Image"""
    try:
        # Remove the data URL prefix if present
        if base64_string.startswith("data:image"):
            base64_string = base64_string.split(",")[1]

        image_data = base64.b64decode(base64_string)
        image = Image.open(io.BytesIO(image_data))
        return image
    except Exception as e:
        print(f"Error decoding base64 image: {e}")
        return None


def process_document_file(uploaded_file):
    """Process document files (PDF, HTML) and convert to markdown using markitdown"""
    try:
        # Initialize markitdown processor
        md = MarkItDown()
        # Convert document to markdown
        result = md.convert(uploaded_file)
        return result.text_content
    except Exception as e:
        print(f"Error processing document: {e}")
        return f"Error processing document file: {str(e)}"


def process_xlsx_file(uploaded_file):
    """Process XLSX file and convert to markdown format"""
    try:
        # Load the workbook
        workbook = load_workbook(uploaded_file)
        markdown_content = ""

        # Process each sheet
        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
            markdown_content += f"\n## Sheet: {sheet_name}\n\n"

            # Build markdown table
            rows = []
            for row in worksheet.iter_rows(values_only=True):
                rows.append(row)

            if rows:
                # Create header
                header = rows[0]
                markdown_content += "| " + " | ".join(str(h) if h is not None else "" for h in header) + " |\n"
                markdown_content += "|" + "|".join(["---"] * len(header)) + "|\n"

                # Add data rows
                for row in rows[1:]:
                    markdown_content += "| " + " | ".join(str(cell) if cell is not None else "" for cell in row) + " |\n"

        return markdown_content
    except Exception as e:
        print(f"Error processing XLSX file: {e}")
        return f"Error processing XLSX file: {str(e)}"


def process_csv_file(uploaded_file):
    """Process CSV file and convert to markdown format"""
    try:
        # Read CSV content
        csv_content = uploaded_file.read().decode('utf-8')
        csv_reader = csv.reader(csv_content.splitlines())

        # Convert to list for easier processing
        rows = list(csv_reader)

        if not rows:
            return "CSV file is empty."

        # Build markdown table
        markdown_content = f"\n## CSV File: {uploaded_file.name}\n\n"

        # Create header
        header = rows[0]
        markdown_content += "| " + " | ".join(str(h) for h in header) + " |\n"
        markdown_content += "|" + "|".join(["---"] * len(header)) + "|\n"

        # Add data rows
        for row in rows[1:]:
            markdown_content += "| " + " | ".join(str(cell) for cell in row) + " |\n"

        return markdown_content
    except Exception as e:
        print(f"Error processing CSV file: {e}")
        return f"Error processing CSV file: {str(e)}"


def process_txt_file(uploaded_file):
    """Process TXT file and convert to markdown format"""
    try:
        # Read text content
        text_content = uploaded_file.read().decode('utf-8')

        # Convert to markdown format
        markdown_content = f"\n## Text File: {uploaded_file.name}\n\n```\n{text_content}\n```\n"

        return markdown_content
    except Exception as e:
        print(f"Error processing TXT file: {e}")
        return f"Error processing TXT file: {str(e)}"


def process_docx_file(uploaded_file):
    """Process DOCX file and convert to markdown format"""
    try:
        # Load the document
        doc = Document(uploaded_file)

        # Extract text from paragraphs
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                full_text.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))

        # Combine all text
        combined_text = "\n\n".join(full_text)

        # Convert to markdown format
        markdown_content = f"\n## Word Document: {uploaded_file.name}\n\n{combined_text}\n"

        return markdown_content
    except Exception as e:
        print(f"Error processing DOCX file: {e}")
        return f"Error processing DOCX file: {str(e)}"


def prepare_chat_history_for_saving(chat_history):
    """Convert chat history to JSON-serializable format by converting PIL Images to base64"""
    serializable_history = []
    for msg in chat_history:
        msg_copy = msg.copy()
        if msg_copy.get("files"):
            # Convert files to serializable format
            serializable_files = []
            for file_data in msg_copy["files"]:
                serializable_file = {
                    "name": file_data["name"],
                    "type": file_data["type"],
                    "processed_type": file_data["processed_type"]
                }
                if file_data["processed_type"] == "image":
                    # Convert PIL Image to base64
                    img_base64 = encode_image_to_base64(file_data["content"])
                    serializable_file["content"] = img_base64
                elif file_data["processed_type"] == "document_markdown":
                    # PDF/HTML/MD content is already text
                    serializable_file["content"] = file_data["content"]
                serializable_files.append(serializable_file)
            msg_copy["files"] = serializable_files
        # Handle legacy images for backward compatibility
        elif msg_copy.get("images"):
            # Convert PIL Images to base64 strings
            images_base64 = []
            for img in msg_copy["images"]:
                if isinstance(img, Image.Image):
                    img_base64 = encode_image_to_base64(img)
                    images_base64.append(img_base64)
                else:
                    # Already base64 string
                    images_base64.append(img)
            msg_copy["images"] = images_base64
        serializable_history.append(msg_copy)
    return serializable_history


def prepare_loaded_chat_history(chat_history):
    """Convert loaded chat history by converting base64 strings back to PIL Images"""
    loaded_history = []
    for msg in chat_history:
        msg_copy = msg.copy()
        if msg_copy.get("files"):
            # Convert loaded files back to working format
            loaded_files = []
            for file_data in msg_copy["files"]:
                loaded_file = {
                    "name": file_data["name"],
                    "type": file_data["type"],
                    "processed_type": file_data["processed_type"]
                }
                if file_data["processed_type"] == "image":
                    # Convert base64 string back to PIL Image
                    img = decode_base64_to_image(file_data["content"])
                    if img:
                        loaded_file["content"] = img
                elif file_data["processed_type"] == "document_markdown":
                    # PDF/HTML/MD content is already text
                    loaded_file["content"] = file_data["content"]
                loaded_files.append(loaded_file)
            msg_copy["files"] = loaded_files
        # Handle legacy images for backward compatibility
        elif msg_copy.get("images"):
            # Convert base64 strings back to PIL Images
            images_pil = []
            for img_base64 in msg_copy["images"]:
                if isinstance(img_base64, str):
                    img = decode_base64_to_image(img_base64)
                    if img:
                        images_pil.append(img)
                else:
                    # Already PIL Image
                    images_pil.append(img_base64)
            msg_copy["images"] = images_pil
        loaded_history.append(msg_copy)
    return loaded_history


def format_message_with_files(text, files):
    """Format message with files (images and PDFs) for API request"""
    content = [{"type": "text", "text": text}]

    # Add file content to text
    file_texts = []
    for file_data in files:
        if file_data["processed_type"] == "document_markdown":
            file_texts.append(f"\n\n--- Document: {file_data['name']} ---\n{file_data['content']}\n--- End Document ---")
        elif file_data["processed_type"] == "image":
            img_base64 = encode_image_to_base64(file_data["content"])
            content.append(
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:image/png;base64,{img_base64}"},
                }
            )

    # Add PDF content as text
    if file_texts:
        content[0]["text"] = text + "".join(file_texts)

    return content


def stream_chat_response(client, model_name, messages, temperature=1.0, extra_body=None):
    """Stream chat response from the selected model"""
    try:
        stream = client.chat.completions.create(
            model=model_name, messages=messages, temperature=temperature, stream=True, extra_body=extra_body
        )

        for chunk in stream:
            if chunk.choices[0].delta.content:
                yield chunk.choices[0].delta.content
    except Exception as e:
        yield f"Error: {str(e)}"


def generate_image_response(model_config, prompt):
    """Generate image using Qwen-Image (ModelScope), Z-Image-Turbo, or Gemini (OpenRouter)"""
    try:
        # Check if it's a ModelScope image model (Qwen-Image or Z-Image-Turbo)
        modelscope_models = ["Qwen/Qwen-Image", "Tongyi-MAI/Z-Image-Turbo"]
        is_modelscope = any(model in model_config["models"] for model in modelscope_models)
        
        if is_modelscope:
            # Extract text prompt if it's a list
            if isinstance(prompt, list):
                for item in prompt:
                    if item["type"] == "text":
                        prompt = item["text"]
                        break

            headers = {
                "Authorization": f"Bearer {model_config['api_key']}",
                "Content-Type": "application/json",
                "X-ModelScope-Async-Mode": "true"
            }
            
            # Initial request
            response = requests.post(
                f"{model_config['base_url']}/images/generations",
                headers=headers,
                data=json.dumps({
                    "model": model_config["models"],
                    "prompt": prompt
                }, ensure_ascii=False).encode('utf-8'),
                timeout=30
            )
            response.raise_for_status()
            task_data = response.json()
            task_id = task_data.get("task_id")
            
            if not task_id:
                raise Exception("No task_id returned from ModelScope API")
            
            print(f"DEBUG: Task created - ID: {task_id}, Initial status: {task_data.get('task_status')}")
            
            # Polling loop with longer initial wait for Z-Image-Turbo
            max_retries = 120
            stuck_counter = 0
            last_status = None
            
            for attempt in range(max_retries):
                time.sleep(2)
                result = requests.get(
                    f"{model_config['base_url']}/tasks/{task_id}",
                    headers={**headers, "X-ModelScope-Task-Type": "image_generation"},
                    timeout=30
                )
                result.raise_for_status()
                data = result.json()
                status = data.get("task_status")
                
                # Track if status is stuck
                if status == last_status:
                    stuck_counter += 1
                else:
                    stuck_counter = 0
                last_status = status

                if attempt % 10 == 0:  # Log every 10 attempts
                    print(f"DEBUG: Poll {attempt+1} - Status: {status}")

                if status == "SUCCEED":
                    output_images = data.get("output_images")
                    if output_images:
                        image_url = output_images[0]
                        print(f"DEBUG: Got image URL: {image_url[:80]}")
                        image_response = requests.get(image_url, timeout=30)
                        image = Image.open(BytesIO(image_response.content))
                        return image
                    else:
                        # Check nested outputs
                        outputs = data.get("outputs", {})
                        if isinstance(outputs, dict) and outputs.get("images"):
                            image_url = outputs["images"][0]
                            image_response = requests.get(image_url, timeout=30)
                            image = Image.open(BytesIO(image_response.content))
                            return image
                        raise Exception("Task completed but no images returned")
                elif status == "FAILED":
                    error = data.get("error", data.get("errors", data.get("message", "Unknown error")))
                    print(f"DEBUG: Task FAILED - Full response: {json.dumps(data, indent=2)[:1000]}")
                    raise Exception(f"ModelScope API error: {error}")
                
                # If stuck in same status for too long, give up
                if stuck_counter > 50:  # Stuck for 100+ seconds
                    if status == "PROCESSING":
                        raise Exception("Image generation stuck in PROCESSING state. The ModelScope API may be experiencing issues. Please try again or use a different model (Gemini Image).")
                    raise Exception(f"Image generation stuck in {status} state for too long")
            
            raise Exception(f"Image generation timeout after {max_retries * 2} seconds. The ModelScope service may be overloaded or unavailable.")

        # Handle OpenAI-compatible image generation (e.g., Gemini via OpenRouter)
        else:
            # Use raw requests for OpenRouter to properly handle the images field
            headers = {
                "Authorization": f"Bearer {model_config['api_key']}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "model": model_config["models"],
                "messages": [
                    {
                        "role": "user",
                        "content": prompt
                    }
                ]
            }
            
            response = requests.post(
                f"{model_config['base_url']}/chat/completions",
                headers=headers,
                json=payload,
                timeout=120
            )
            
            if response.status_code == 402:
                raise Exception("Payment Required: Please check your OpenRouter credits or use a free model.")
            
            response.raise_for_status()
            result = response.json()
            
            # Debug: print the response
            print(f"DEBUG - Response JSON: {json.dumps(result, indent=2)[:1000]}")
            
            # Check for images in the response
            if result.get("choices"):
                message = result["choices"][0]["message"]
                
                # Check for images field (OpenRouter Gemini format)
                if message.get("images"):
                    for image_data in message["images"]:
                        image_url = image_data["image_url"]["url"]
                        if image_url.startswith("data:image"):
                            image = decode_base64_to_image(image_url)
                            return image
                        else:
                            image_response = requests.get(image_url)
                            image = Image.open(BytesIO(image_response.content))
                            return image
                
            # Check if content has image data
                content = message.get("content", "")
                if content:
                    # Check for base64 image in content string
                    import re
                    match = re.search(r"data:image/[a-zA-Z]+;base64,[a-zA-Z0-9+/=]+", content)
                    
                    if match:
                        image_data_url = match.group(0)
                        image = decode_base64_to_image(image_data_url)
                        return image
                    
                    # Check for URL
                    url_match = re.search(r"https?://[^\s\)]+", content)
                    if url_match:
                        image_url = url_match.group(0).strip(")")
                        image_response = requests.get(image_url)
                        image = Image.open(BytesIO(image_response.content))
                        return image
            
            error_msg = "No image found in response."
            if content:
                error_msg += f" Model responded: {content}"
            raise Exception(error_msg)
            
    except Exception as e:
        if "402" in str(e):
             raise Exception("Payment Required: Please check your OpenRouter credits.")
        raise Exception(f"Image generation error: {str(e)}")


import tempfile
from dashscope.files import Files

def upload_image_to_oss(image, api_key):
    """Upload image to DashScope OSS and return URL using direct API"""
    try:
        # Create a temporary file
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_file:
            image.save(temp_file.name, format="PNG")
            temp_path = temp_file.name

        # Upload to DashScope via REST API
        # Use the provided API key or fallback to loaded keys
        api_keys = load_api_keys()
        upload_api_key = api_key if api_key else api_keys.get("dashscope", "")
        
        base_url = "https://dashscope.aliyuncs.com/api/v1/files"
        headers = {"Authorization": f"Bearer {upload_api_key}"}
        
        try:
            # Step 1: Upload file
            with open(temp_path, 'rb') as f:
                files = {'file': ('image.png', f, 'image/png')}
                data = {'purpose': 'general'}
                
                response = requests.post(base_url, headers=headers, files=files, data=data)
                
                if response.status_code != 200:
                    raise Exception(f"Upload failed ({response.status_code}): {response.text}")
                
                resp_json = response.json()
                file_id = resp_json['data']['uploaded_files'][0]['file_id']
                
            # Step 2: Get file URL
            detail_url = f"{base_url}/{file_id}"
            response = requests.get(detail_url, headers=headers)
            
            if response.status_code != 200:
                raise Exception(f"Get details failed ({response.status_code}): {response.text}")
                
            resp_json = response.json()
            if 'url' in resp_json['data']:
                return resp_json['data']['url']
            else:
                raise Exception(f"No URL in file details: {resp_json}")
                
        except Exception as upload_error:
            raise Exception(f"DashScope Upload failed: {str(upload_error)}")
            
        finally:
            # Clean up temp file
            if os.path.exists(temp_path):
                os.remove(temp_path)

    except Exception as e:
        raise Exception(f"Error preparing image for upload: {str(e)}")


def generate_image_edit_response(model_config, prompt, image_file):
    """Generate image edit response using Qwen-Image-Edit-Plus or Flux"""
    try:
        # Check if it's a Flux model (ModelScope)
        if "flux" in model_config["models"].lower():
            # Flux requires a public HTTP URL, not base64
            # We use DashScope's file upload service to get a temp URL
            # Note: DashScope API key should be available from .env
            
            print("DEBUG: Uploading image for Flux...")
            api_keys = load_api_keys()
            image_url = upload_image_to_oss(image_file, api_keys.get("dashscope", ""))
            print(f"DEBUG: Image uploaded to: {image_url}")
            
            headers = {
                "Authorization": f"Bearer {model_config['api_key']}",
                "Content-Type": "application/json",
                "X-ModelScope-Async-Mode": "true"
            }
            
            payload = {
                "model": model_config["models"],
                "prompt": prompt,
                "image_url": [image_url]  # Use the uploaded URL
            }
            
            # Initial request
            response = requests.post(
                f"{model_config['base_url']}/images/generations",
                headers=headers,
                data=json.dumps(payload, ensure_ascii=False).encode('utf-8'),
                timeout=30
            )
            
            if response.status_code != 200:
                print(f"DEBUG: ModelScope Error Response: {response.text}")
                raise Exception(f"ModelScope API Error ({response.status_code}): {response.text}")

            task_data = response.json()
            task_id = task_data.get("task_id")
            
            if not task_id:
                raise Exception("No task_id returned from ModelScope API")
            
            # Polling loop
            max_retries = 60
            for attempt in range(max_retries):
                time.sleep(2)
                result = requests.get(
                    f"{model_config['base_url']}/tasks/{task_id}",
                    headers={**headers, "X-ModelScope-Task-Type": "image_generation"},
                    timeout=30
                )
                result.raise_for_status()
                data = result.json()
                status = data.get("task_status")
                
                if status == "SUCCEED":
                    output_images = data.get("output_images")
                    if output_images:
                        out_image_url = output_images[0]
                        image_response = requests.get(out_image_url, timeout=30)
                        image = Image.open(BytesIO(image_response.content))
                        return image
                    raise Exception("Task completed but no images returned")
                elif status == "FAILED":
                    error = data.get("error", data.get("errors", data.get("message", "Unknown error")))
                    raise Exception(f"ModelScope API error: {error}")
                
            raise Exception("Image editing timeout")

        # Fallback to DashScope SDK for Qwen
        else:
            # Encode image to data URI format required by DashScope
            buffered = io.BytesIO()
            image_file.save(buffered, format="PNG")
            encoded_string = base64.b64encode(buffered.getvalue()).decode('utf-8')
            image_data = f"data:image/png;base64,{encoded_string}"

            messages = [
                {
                    "role": "user",
                    "content": [
                        {"image": image_data},
                        {"text": prompt}
                    ]
                }
            ]

            response = MultiModalConversation.call(
                api_key=model_config["api_key"],
                model=model_config["models"],
                messages=messages,
                stream=False,
                n=1, # Generate 1 image
            )

            if response.status_code == 200:
                image_url = response.output.choices[0].message.content[0]['image']
                image_response = requests.get(image_url)
                image = Image.open(BytesIO(image_response.content))
                return image
            else:
                raise Exception(f"Error: {response.code} - {response.message}")

    except Exception as e:
        raise Exception(f"Image edit error: {str(e)}")


def stream_model_response(
    model_config, messages, temperature, response_queue, model_name
):
    """Helper function to stream response from a single model and put chunks in queue"""
    try:
        # Handle Image Generation Models
        if model_config.get("type") == "image":
            # Get the last user message as prompt
            # Get the last user message as prompt
            prompt = messages[-1]["content"]
            # Pass full prompt (text or list) to generate_image_response
            
            response_queue.put(("chunk", model_name, "🎨 Generating image...", "Generating..."))
            
            try:
                image = generate_image_response(model_config, prompt)
                response_queue.put(("image", model_name, image))
                response_queue.put(("done", model_name, "Image generated successfully"))
            except Exception as e:
                response_queue.put(("error", model_name, str(e)))
            return

        # Handle Image Edit Models
        if model_config.get("type") == "image_edit":
            # Get the last user message as prompt
            prompt = messages[-1]["content"]
            # Extract image from message if available
            # In this app structure, images are attached to the message or uploaded
            # We need to find the image associated with this request.
            # The 'messages' passed here are chat history format.
            # However, for image edit, we expect an image to be uploaded in the current turn.
            
            # Check if there are uploaded files in the current session state that haven't been processed?
            # Or check the last message for image content if it was formatted that way.
            
            # Let's look at how messages are constructed. 
            # format_message_with_files constructs the message.
            # But here 'messages' is the list of messages sent to the API.
            
            # We need the actual image object.
            # In the current flow, images are converted to base64 in 'messages'.
            # But for this specific model, we might want to grab the image from st.session_state.uploaded_files
            # if this is the first turn, or from history if it's a follow up (though this model is single turn).
            
            # Simplest approach: Check the last message content for image_url
            last_message = messages[-1]
            image_to_edit = None
            
            if isinstance(last_message["content"], list):
                for item in last_message["content"]:
                    if item["type"] == "text":
                        prompt = item["text"]
                    elif item["type"] == "image_url":
                        # Decode base64 to PIL Image
                        image_url = item["image_url"]["url"]
                        image_to_edit = decode_base64_to_image(image_url)
            
            if not image_to_edit:
                 response_queue.put(("error", model_name, "Please upload an image to edit."))
                 return

            response_queue.put(("chunk", model_name, "🎨 Editing image...", "Editing..."))
            
            try:
                image = generate_image_edit_response(model_config, prompt, image_to_edit)
                response_queue.put(("image", model_name, image))
                response_queue.put(("done", model_name, "Image edited successfully"))
            except Exception as e:
                response_queue.put(("error", model_name, str(e)))
            return

        client = create_client(model_config)
        if not client:
            response_queue.put(("error", model_name, "Unable to create client"))
            return

        api_model_name = model_config["models"]
        full_response = ""

        # Check for reasoning flag in model config
        extra_body = None
        if model_config.get("reasoning"):
            extra_body = {
                "reasoning": {
                    "effort": "high",
                    "exclude": True
                }
            }

        for chunk in stream_chat_response(
            client, api_model_name, messages, temperature, extra_body
        ):
            full_response += chunk
            response_queue.put(("chunk", model_name, full_response, chunk))

        response_queue.put(("done", model_name, full_response))

    except Exception as e:
        response_queue.put(("error", model_name, str(e)))


def save_current_chat():
    """Save current chat to sessions and file"""
    if st.session_state.chat_history:
        chat_id = st.session_state.current_chat_id
        # Generate a title from first user message
        title = "New Chat"
        for msg in st.session_state.chat_history:
            if msg["role"] == "user":
                title = (
                    msg["content"][:30] + "..."
                    if len(msg["content"]) > 30
                    else msg["content"]
                )
                break

        st.session_state.chat_sessions[chat_id] = {
            "title": title,
            "history": prepare_chat_history_for_saving(st.session_state.chat_history),
            "timestamp": datetime.now().isoformat(),
        }
        # Save to file
        save_chat_history_to_file()


def new_chat():
    """Start a new chat session"""
    save_current_chat()
    st.session_state.chat_history = []
    st.session_state.current_chat_id = f"chat_{int(time.time())}"
    st.session_state.uploaded_files = []
    # Increment file uploader key to clear the widget
    st.session_state.file_uploader_key += 1


def load_chat(chat_id):
    """Load a specific chat session"""
    if chat_id in st.session_state.chat_sessions:
        save_current_chat()  # Save current chat before loading new one
        loaded_history = prepare_loaded_chat_history(st.session_state.chat_sessions[chat_id]["history"])
        st.session_state.chat_history = loaded_history
        st.session_state.current_chat_id = chat_id
        st.session_state.uploaded_files = []
        # Increment file uploader key to clear the widget
        st.session_state.file_uploader_key += 1


def delete_chat(chat_id):
    """Delete a chat session"""
    if chat_id in st.session_state.chat_sessions:
        del st.session_state.chat_sessions[chat_id]
        # Save to file after deletion
        save_chat_history_to_file()
        if st.session_state.current_chat_id == chat_id:
            new_chat()


# Initialize session state (after all functions are defined)
if "chat_sessions" not in st.session_state:
    # Load from file if exists, otherwise empty dict
    st.session_state.chat_sessions = load_chat_history()

# Load settings
settings = load_settings()

# Initialize chat history and current chat ID
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "current_chat_id" not in st.session_state:
    # Always start with a new chat session when app loads
    st.session_state.current_chat_id = f"chat_{int(time.time())}"
if "current_models" not in st.session_state:
    st.session_state.current_models = settings["current_models"]
if "api_key" not in st.session_state:
    st.session_state.api_key = ""
if "base_url" not in st.session_state:
    st.session_state.base_url = ""
if "temperature" not in st.session_state:
    st.session_state.temperature = settings["temperature"]
if "system_prompt" not in st.session_state:
    st.session_state.system_prompt = settings["system_prompt"]
if "system_prompt_widget" not in st.session_state:
    st.session_state.system_prompt_widget = st.session_state.system_prompt
if "uploaded_files" not in st.session_state:
    st.session_state.uploaded_files = []
if "language" not in st.session_state:
    st.session_state.language = settings["language"]
if "file_uploader_key" not in st.session_state:
    st.session_state.file_uploader_key = 0
if "web_search_enabled" not in st.session_state:
    st.session_state.web_search_enabled = settings.get("web_search_enabled", "auto")


# Main application layout
st.set_page_config(
    page_title="AI Chat Application",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Reduce spacing with custom CSS and fix input box at bottom
st.markdown("""
<style>
    .block-container {
        padding-top: 3rem !important;
        padding-bottom: 0rem !important;
        padding-left: 1rem !important;
        padding-right: 1rem !important;
    }
    .css-1d391kg {
        padding-top: 0rem !important;
    }
    .stTitle {
        margin-top: 4.5rem !important;
        margin-bottom: 2rem !important;
    }
    .css-17ziqus {
        padding-top: 0rem !important;
    }
    .css-1lcbmhc {
        padding-top: 0.5rem !important;
    }
    .element-container {
        margin-top: -0.5rem !important;
        margin-bottom: 0.5rem !important;
    }
    .stHeader {
        margin-top: -1rem !important;
        margin-bottom: 0.5rem !important;
    }
    .css-1lcbmhc .element-container {
        margin-top: 0rem !important;
    }
    .css-1d391kg {
        padding-top: 0rem !important;
    }
    .css-1oe5foe {
        padding-top: 0rem !important;
    }
    .css-17eq3hr {
        padding-top: 0rem !important;
    }
    .css-1vq4p4l {
        padding-top: 0rem !important;
    }
    .stSidebar {
        padding-top: 0rem !important;
    }
    /* Chat input styling to keep it at bottom */
    .stChatInput {
        position: fixed !important;
        bottom: 0 !important;
        left: 280px !important;
        right: 0 !important;
        background-color: white !important;
        padding: 1rem !important;
        border-top: 1px solid #f0f0f0 !important;
        z-index: 50 !important;
        box-shadow: 0 -2px 10px rgba(0,0,0,0.1) !important;
    }

    /* For collapsed sidebar */
    @media (max-width: 768px) {
        .stChatInput {
            left: 80px !important;
        }
    }

    /* For mobile/collapsed sidebar */
    @media (max-width: 480px) {
        .stChatInput {
            left: 0px !important;
        }
    }
    /* Add padding to main content to avoid overlap with fixed input */
    .main .block-container {
        padding-bottom: 500px !important;
    }
    /* Style for chat tab specifically - removed background color */
    .stTabs [data-baseweb="tab-list"] {
        margin-bottom: 1rem !important;
    }
    .stTabs [data-baseweb="tab"] {
        height: 50px !important;
    }
</style>
""", unsafe_allow_html=True)

st.title(t("app_title"))

# SIDEBAR - Configuration with Tabs
with st.sidebar:
    st.markdown(f"### {t('config')}")

    # Create tabs
    tab1, tab2, tab3 = st.tabs(["🤖 Models", "⚙️ Parameters", t("keys")])

    # Tab 1: Models
    with tab1:
        # Model selection - Multiple choice
        available_models = get_all_models()
        
        # Sort models by capability
        text_to_text = []
        multimodal_vision = []
        text_to_image = []
        image_edit = []
        
        for model_name, model_config in available_models.items():
            input_type = model_config.get("input_type", {"text"})
            output_type = model_config.get("output_type", {"text"})
            
            # Convert to set if not already
            if not isinstance(input_type, set):
                input_type = {input_type}
            if not isinstance(output_type, set):
                output_type = {output_type}
            
            if "image" in output_type and "image" in input_type:
                image_edit.append(model_name)
            elif "image" in output_type:
                text_to_image.append(model_name)
            elif "image" in input_type:
                multimodal_vision.append(model_name)
            else:
                text_to_text.append(model_name)
                
        # Combine lists in desired order
        model_names = text_to_text + multimodal_vision + text_to_image + image_edit

        def on_model_change():
            """Callback to save settings when model selection changes"""
            st.session_state.current_models = st.session_state.model_selector
            save_settings()

        # Filter current_models to only include available models
        available_current_models = [
            model for model in st.session_state.current_models
            if model in model_names
        ]

        selected_model_names = st.multiselect(
            t("select_models"),
            options=model_names,
            default=available_current_models,
            key="model_selector",
            on_change=on_model_change
        )

        # Keep current_models in sync (for backward compatibility)
        if "model_selector" in st.session_state:
            st.session_state.current_models = st.session_state.model_selector

        # Add new model button
        if st.button("➕ Add New Model", use_container_width=True):
            if "show_add_model_form" not in st.session_state:
                st.session_state.show_add_model_form = False
            st.session_state.show_add_model_form = not st.session_state.show_add_model_form
            st.rerun()

        # Show add model form if toggled
        if st.session_state.get("show_add_model_form", False):
            st.markdown("### Add New Model")
            with st.form("add_model_form"):
                new_model_name = st.text_input("Model Name", placeholder="e.g., Custom GPT-4")
                new_api_key = st.text_input("API Key", type="password", placeholder="Enter your API key")
                new_base_url = st.text_input("Base URL", placeholder="e.g., https://api.openai.com/v1")
                new_models = st.text_input("Model", placeholder="e.g., gpt-4", help="input model name")

                col1, col2 = st.columns(2)
                with col1:
                    submit_button = st.form_submit_button("Add Model", use_container_width=True)
                with col2:
                    cancel_button = st.form_submit_button("Cancel", use_container_width=True)

                if submit_button:
                    if new_model_name and new_api_key and new_base_url and new_models:
                        add_custom_model(new_model_name, new_api_key, new_base_url, new_models)
                        st.session_state.show_add_model_form = False
                        # Auto-select the newly added model
                        if new_model_name not in st.session_state.current_models:
                            st.session_state.current_models.append(new_model_name)
                        st.success(f"Model '{new_model_name}' added successfully!")
                        st.rerun()
                    else:
                        st.error("Please fill in all fields")

                if cancel_button:
                    st.session_state.show_add_model_form = False
                    st.rerun()

        # File upload section
        st.subheader(t("upload_images"))
        uploaded_files = st.file_uploader(
            t("choose_images"),
            type=["png", "jpg", "jpeg", "pdf", "html", "htm", "md", "markdown", "xlsx", "csv", "txt", "docx"],
            accept_multiple_files=True,
            help=t("upload_help"),
            label_visibility="collapsed",
            key=f"file_uploader_{st.session_state.file_uploader_key}"
        )

        # Process uploaded files
        if uploaded_files:
            st.session_state.uploaded_files = []
            for uploaded_file in uploaded_files:
                file_data = {
                    "name": uploaded_file.name,
                    "type": uploaded_file.type,
                    "content": None,
                    "file": uploaded_file
                }

                if uploaded_file.type == "application/pdf":
                    # Process PDF
                    document_content = process_document_file(uploaded_file)
                    file_data["content"] = document_content
                    file_data["processed_type"] = "document_markdown"
                elif uploaded_file.type == "text/html" or uploaded_file.name.lower().endswith((".html", ".htm")):
                    # Process HTML
                    document_content = process_document_file(uploaded_file)
                    file_data["content"] = document_content
                    file_data["processed_type"] = "document_markdown"
                elif uploaded_file.type == "text/markdown" or uploaded_file.name.lower().endswith((".md", ".markdown")):
                    # Process Markdown - read as text
                    markdown_content = uploaded_file.read().decode("utf-8")
                    file_data["content"] = markdown_content
                    file_data["processed_type"] = "document_markdown"
                elif uploaded_file.name.lower().endswith(".xlsx") or "spreadsheet" in uploaded_file.type:
                    # Process XLSX
                    xlsx_content = process_xlsx_file(uploaded_file)
                    file_data["content"] = xlsx_content
                    file_data["processed_type"] = "document_markdown"
                elif uploaded_file.name.lower().endswith(".csv") or "csv" in uploaded_file.type:
                    # Process CSV
                    csv_content = process_csv_file(uploaded_file)
                    file_data["content"] = csv_content
                    file_data["processed_type"] = "document_markdown"
                elif uploaded_file.name.lower().endswith(".txt") or "text" in uploaded_file.type:
                    # Process TXT
                    txt_content = process_txt_file(uploaded_file)
                    file_data["content"] = txt_content
                    file_data["processed_type"] = "document_markdown"
                elif uploaded_file.name.lower().endswith(".docx") or "word" in uploaded_file.type or "openxmlformats-officedocument" in uploaded_file.type:
                    # Process DOCX
                    docx_content = process_docx_file(uploaded_file)
                    file_data["content"] = docx_content
                    file_data["processed_type"] = "document_markdown"
                else:
                    # Process image
                    image = Image.open(uploaded_file)
                    file_data["content"] = image
                    file_data["processed_type"] = "image"

                st.session_state.uploaded_files.append(file_data)

            # Display uploaded files
            st.write(t("uploaded_images_title"))
            cols = st.columns(min(len(st.session_state.uploaded_files), 2))
            for i, file_data in enumerate(st.session_state.uploaded_files):
                with cols[i % 2]:
                    if file_data["processed_type"] == "image":
                        st.image(file_data["content"], width=150, caption=f"Image: {file_data['name']}")
                    elif file_data["processed_type"] == "document_markdown":
                        # Display appropriate icon based on file extension
                        if file_data['name'].lower().endswith('.pdf'):
                            st.info(f"📄 {file_data['name']} (PDF)")
                        elif file_data['name'].lower().endswith('.xlsx'):
                            st.info(f"📊 {file_data['name']} (Excel)")
                        elif file_data['name'].lower().endswith(('.html', '.htm')):
                            st.info(f"🌐 {file_data['name']} (HTML)")
                        elif file_data['name'].lower().endswith(('.md', '.markdown')):
                            st.info(f"📝 {file_data['name']} (Markdown)")
                        elif file_data['name'].lower().endswith('.csv'):
                            st.info(f"📊 {file_data['name']} (CSV)")
                        elif file_data['name'].lower().endswith('.txt'):
                            st.info(f"📄 {file_data['name']} (Text)")
                        elif file_data['name'].lower().endswith('.docx'):
                            st.info(f"📘 {file_data['name']} (Word)")
                        else:
                            st.info(f"📄 {file_data['name']} (Document)")

        # Display current model info
        # st.info(f"**Model:** {selected_model_name}\n**Base URL:** {current_model_config['base_url']}")

        # API Key input (masked for security)
        # api_key_display = st.session_state.api_key if st.session_state.api_key else current_model_config.get("api_key", "")
        # if api_key_display:
        #     masked_key = api_key_display[:8] + "..." if len(api_key_display) > 8 else "***"
        #     st.text_input("API Key", value=masked_key, disabled=True, type="password")
        # else:
        #     st.warning("⚠️ API Key not configured")

        # Chat History Section (in Models tab)
        st.divider()
        st.subheader(t("chat_history"))

        # New Chat button
        if st.button(t("new_chat"), use_container_width=True, type="primary", key="new_chat_models"):
            new_chat()
            st.rerun()

        # Display saved chats
        if st.session_state.chat_sessions:
            st.write(t("saved_chats"))
            for chat_id, chat_data in st.session_state.chat_sessions.items():
                col1, col2 = st.columns([4, 1])
                with col1:
                    if st.button(
                        chat_data["title"],
                        key=f"load_{chat_id}",
                        use_container_width=True,
                        help=chat_data["timestamp"],
                    ):
                        load_chat(chat_id)
                        st.rerun()
                with col2:
                    if st.button("🗑️", key=f"delete_{chat_id}", help=t("delete_chat")):
                        delete_chat(chat_id)
                        st.rerun()
        else:
            st.write(t("no_saved_chats"))

    # Tab 2: Parameters
    with tab2:
        # Parameters
        st.subheader(t("parameters"))

        def on_temperature_change():
            """Callback to save temperature when it changes"""
            # Update session state with new temperature from widget
            st.session_state.temperature = st.session_state.temperature_widget
            save_settings()

        temperature = st.slider(
            t("temperature"),
            min_value=0.0,
            max_value=1.5,
            value=float(st.session_state.temperature) if isinstance(st.session_state.temperature, (int, float, str)) else 1.0,
            step=0.1,
            help=t("temperature_help"),
            key="temperature_widget",
            on_change=on_temperature_change
        )

        def on_system_prompt_change():
            """Callback to save system prompt when it changes"""
            st.session_state.system_prompt = st.session_state.system_prompt_widget
            save_settings()

        system_prompt = st.text_area(
            t("system_prompt"),
            height=300,
            help=t("system_prompt_help"),
            key="system_prompt_widget",
            on_change=on_system_prompt_change
        )

        # Web Search Settings
        st.divider()
        st.subheader("🔍 Web Search Settings")

        def on_web_search_change():
            """Callback to save web search setting when it changes"""
            st.session_state.web_search_enabled = st.session_state.web_search_enabled_widget
            save_settings()

        # Create options for search mode
        search_options = {
            "off": t("search_mode_off"),
            "on": t("search_mode_always")
        }

        # Get current index safely
        current_mode = st.session_state.web_search_enabled
        try:
            current_index = list(search_options.keys()).index(current_mode)
        except ValueError:
            current_index = list(search_options.keys()).index("off")  # Default to off

        web_search_mode = st.selectbox(
            t("web_search"),
            options=search_options,
            index=current_index,
            help=t("web_search_help"),
            key="web_search_enabled_widget",
            on_change=on_web_search_change
        )

        # Show search API key status if search is on
        if web_search_mode == "on":
            try:
                search_manager = get_search_manager()
                st.success("✅ Tavily API key is configured")
            except Exception as e:
                st.error("❌ Tavily API key not found or invalid")
                st.info("Please add 'tavily_api_key=your_api_key' to your .env file")

        # Model Info content moved here
        st.divider()
        st.subheader("📊 Model Capabilities")

        # Get all models and categorize them
        available_models = get_all_models()

        # Categorize models by type
        text_to_text = []
        multimodal_vision = []
        text_to_image = []
        image_edit = []

        for model_name, model_config in available_models.items():
            input_type = model_config.get("input_type", {"text"})
            output_type = model_config.get("output_type", {"text"})

            # Convert to set if not already
            if not isinstance(input_type, set):
                input_type = {input_type}
            if not isinstance(output_type, set):
                output_type = {output_type}

            if "image" in output_type and "image" in input_type:
                # Image editing models
                image_edit.append(model_name)
            elif "image" in output_type:
                # Image generation models
                text_to_image.append(model_name)
            elif "image" in input_type:
                # Multimodal vision models
                multimodal_vision.append(model_name)
            else:
                # Pure text models
                text_to_text.append(model_name)

        # Display categorized models
        st.markdown("### Pure Text Models")
        if text_to_text:
            st.markdown(f"**Text → Text:** {', '.join(text_to_text)}")
        else:
            st.info("No text-only models available")

        st.markdown("### Multimodal Vision Models")
        if multimodal_vision:
            st.markdown(f"**Text+Image → Text:** {', '.join(multimodal_vision)}")
        else:
            st.info("No multimodal vision models available")

        st.markdown("### Image Generation Models")
        if text_to_image:
            st.markdown(f"**Text → Image:** {', '.join(text_to_image)}")

        if image_edit:
            st.markdown(f"**Text+Image → Image:** {', '.join(image_edit)} (image editing)")

        if not text_to_image and not image_edit:
             st.info("No image generation models available")

    # Tab 3: Keys
    with tab3:
        st.subheader(t("api_keys"))
        st.caption(t("api_keys_help"))

        # Load current API keys
        current_keys = load_api_keys()

        # Check if .env file exists
        env_exists = os.path.exists(ENV_FILE)
        if env_exists:
            st.info(t("keys_env_file"))
        else:
            st.info("💡 Enter your API keys below. They will be saved automatically.")

        # Function to save a single API key
        def save_single_key(key_name, widget_key):
            """Save a single API key when it changes"""
            new_value = st.session_state.get(widget_key, "").strip()
            current = load_api_keys()
            if current.get(key_name) != new_value:
                current[key_name] = new_value
                if save_api_keys_to_env(current):
                    # Reload config module to pick up new keys
                    import importlib
                    importlib.reload(config)
                    if 'search_providers' in sys.modules:
                        importlib.reload(sys.modules['search_providers'])

        # API Key inputs with auto-save on change
        # OpenRouter Key
        st.text_input(
            t("openrouter_key"),
            value=current_keys.get("openrouter", ""),
            type="password",
            help="API key for OpenRouter services",
            key="openrouter_key_input",
            on_change=save_single_key,
            args=("openrouter", "openrouter_key_input")
        )

        # ModelScope Key
        st.text_input(
            t("modelscope_key"),
            value=current_keys.get("modelscope", ""),
            type="password",
            help="API key for ModelScope services",
            key="modelscope_key_input",
            on_change=save_single_key,
            args=("modelscope", "modelscope_key_input")
        )

        # SiliconFlow Key
        st.text_input(
            t("siliconflow_key"),
            value=current_keys.get("siliconflow", ""),
            type="password",
            help="API key for SiliconFlow services",
            key="siliconflow_key_input",
            on_change=save_single_key,
            args=("siliconflow", "siliconflow_key_input")
        )

        # DashScope Key
        st.text_input(
            t("dashscope_key"),
            value=current_keys.get("dashscope", ""),
            type="password",
            help="API key for DashScope services",
            key="dashscope_key_input",
            on_change=save_single_key,
            args=("dashscope", "dashscope_key_input")
        )

        # BigModel Key
        st.text_input(
            t("bigmodel_key"),
            value=current_keys.get("bigmodel", ""),
            type="password",
            help="API key for BigModel services",
            key="bigmodel_key_input",
            on_change=save_single_key,
            args=("bigmodel", "bigmodel_key_input")
        )

        # Tavily Search Key
        st.text_input(
            t("tavily_key"),
            value=current_keys.get("tavily_api_key", ""),
            type="password",
            help="API key for Tavily web search",
            key="tavily_key_input",
            on_change=save_single_key,
            args=("tavily_api_key", "tavily_key_input")
        )

        # Test Keys button (optional, outside of form now)
        st.divider()
        if st.button("🧪 Test Keys", use_container_width=True):
            # Reload current keys
            test_keys = load_api_keys()
            with st.spinner("Testing API keys..."):
                # Test Tavily key (web search)
                if test_keys.get("tavily_api_key", "").strip():
                    try:
                        from search_providers import get_search_manager
                        os.environ['tavily_api_key'] = test_keys["tavily_api_key"].strip()
                        search_manager = get_search_manager()
                        import asyncio
                        result = asyncio.run(search_manager.perform_search("test"))
                        if result:
                            st.success("✅ Tavily API key is valid")
                    except Exception as e:
                        st.error(f"❌ Tavily API key error: {str(e)}")

                # Show status for all keys
                key_status = {
                    "openrouter": "✅" if test_keys.get("openrouter", "").strip() else "⚠️ Not set",
                    "modelscope": "✅" if test_keys.get("modelscope", "").strip() else "⚠️ Not set",
                    "siliconflow": "✅" if test_keys.get("siliconflow", "").strip() else "⚠️ Not set",
                    "dashscope": "✅" if test_keys.get("dashscope", "").strip() else "⚠️ Not set",
                    "bigmodel": "✅" if test_keys.get("bigmodel", "").strip() else "⚠️ Not set",
                    "tavily_api_key": "✅" if test_keys.get("tavily_api_key", "").strip() else "⚠️ Not set"
                }

                st.write("**Key Status:**")
                for key_name, status in key_status.items():
                    st.write(f"- {key_name}: {status}")



# MAIN PANEL - Chat Interface
# Language switch in top right corner
col1, col2 = st.columns([11, 1])
with col1:
    #st.markdown(f"### {t('chat_interface')}")
    st.markdown("")
with col2:
    current_lang = st.session_state.language
    next_lang = "中文" if current_lang == "EN" else "EN"
    if st.button(next_lang, use_container_width=True, key="lang_switch_top"):
        st.session_state.language = next_lang
        save_settings()
        st.rerun()

# Main Tabs
tab_chat, tab_prompt_bay = st.tabs(["💬 Chat", "📚 Prompt Bay"])

with tab_prompt_bay:
    st.subheader("Prompt Bay")
    st.markdown("Select a system prompt from the library below.")
    
    try:
        # Check if file exists
        if os.path.exists("prompt_bay.csv"):
            df = pd.read_csv("prompt_bay.csv")
            
            # Check required columns
            if 'act' in df.columns and 'prompt' in df.columns:
                # Search filter
                search_query = st.text_input("Search prompts...", key="prompt_search")
                if search_query:
                    mask = df['act'].astype(str).str.contains(search_query, case=False) | df['prompt'].astype(str).str.contains(search_query, case=False)
                    df = df[mask]
                
                def apply_prompt_callback(prompt_text):
                    """Callback to update system prompt state"""
                    st.session_state.system_prompt = prompt_text
                    st.session_state.system_prompt_widget = prompt_text
                    save_settings()

                for index, row in df.iterrows():
                    with st.container():
                        col_p1, col_p2 = st.columns([5, 1])
                        with col_p1:
                            st.markdown(f"**{row['act']}**")
                            # Truncate prompt for display
                            prompt_text = str(row['prompt'])
                            display_prompt = prompt_text[:200] + "..." if len(prompt_text) > 200 else prompt_text
                            st.caption(display_prompt)
                        with col_p2:
                            if st.button("Apply", key=f"apply_{index}", on_click=apply_prompt_callback, args=(row['prompt'],)):
                                st.toast(f"Applied prompt: {row['act']}")
                    st.divider()
            else:
                st.error("CSV file must contain 'act' and 'prompt' columns.")
        else:
            st.info("prompt_bay.csv not found.")
    except Exception as e:
        st.error(f"Error loading prompt bay: {str(e)}")

with tab_chat:
    # Create a container for chat history that will expand
    # and a fixed bottom container for input
    chat_history_container = st.container()

    with chat_history_container:
        # Display chat history
        i = 0
        while i < len(st.session_state.chat_history):
            message = st.session_state.chat_history[i]

            if message["role"] == "user":
                with st.chat_message("user"):
                    st.write(message["content"])

                    # Display search results if available
                    if message.get("search_results"):
                        with st.expander(f"🔍 {t('search_results')}", expanded=False):
                            search_data = message["search_results"]
                            if search_data.get("results"):
                                for j, result in enumerate(search_data["results"], 1):
                                    st.markdown(f"**{j}. {result['title']}**")
                                    st.markdown(f"🔗 [{result['url']}]({result['url']})")
                                    st.markdown(f"📝 {result['snippet'][:300]}{'...' if len(result['snippet']) > 300 else ''}")
                                    if result.get("date"):
                                        st.caption(f"📅 {result['date']} | 🌐 {result.get('source', 'Unknown')}")
                                    st.divider()
                            else:
                                st.info(t("no_search_results"))

                    # Display files if any
                    if message.get("files"):
                        for file_data in message["files"]:
                            if file_data["processed_type"] == "image":
                                st.image(file_data["content"], width=300, caption=f"Image: {file_data['name']}")
                            elif file_data["processed_type"] == "document_markdown":
                                with st.expander(f"📄 {file_data['name']} (Document Content)"):
                                    st.markdown(file_data["content"])
                i += 1
            else:
                # Group consecutive assistant messages (multiple model responses)
                assistant_messages = []
                while (
                    i < len(st.session_state.chat_history)
                    and st.session_state.chat_history[i]["role"] == "assistant"
                ):
                    assistant_messages.append(st.session_state.chat_history[i])
                    i += 1

                # Display assistant messages
                if len(assistant_messages) == 1:
                    # Single message - display normally
                    with st.chat_message("assistant"):
                        content = assistant_messages[0]["content"]
                        st.write(content)
                        # Display files if any (e.g., generated images)
                        if assistant_messages[0].get("files"):
                            for file_data in assistant_messages[0]["files"]:
                                if file_data["processed_type"] == "image":
                                    st.image(file_data["content"], caption=f"Generated Image", width=512)
                                    # Add download button
                                    buf = io.BytesIO()
                                    file_data["content"].save(buf, format="PNG")
                                    byte_im = buf.getvalue()
                                    st.download_button(
                                        label="Download Image",
                                        data=byte_im,
                                        file_name=file_data["name"],
                                        mime="image/png"
                                    )
                else:
                    # Multiple messages - display horizontally in columns
                    cols = st.columns(len(assistant_messages))
                    for j, assistant_msg in enumerate(assistant_messages):
                        with cols[j]:
                            content = assistant_msg["content"]
                            st.write(content)
                            # Display files if any (e.g., generated images)
                            if assistant_msg.get("files"):
                                for file_data in assistant_msg["files"]:
                                    if file_data["processed_type"] == "image":
                                        st.image(file_data["content"], caption=f"Generated Image", width=512)
                                        # Add download button
                                        buf = io.BytesIO()
                                        file_data["content"].save(buf, format="PNG")
                                        byte_im = buf.getvalue()
                                        st.download_button(
                                            label="Download Image",
                                            data=byte_im,
                                            file_name=file_data["name"],
                                            mime="image/png",
                                            key=f"download_history_{j}_{file_data['name']}"
                                        )

        # Spacer to prevent the fixed input bar from covering the last messages
        st.markdown("<div style='height:60px'></div>", unsafe_allow_html=True)

    # Fixed bottom input container without divider
    user_input_container = st.container()
    with user_input_container:
        # User input section - fixed at the bottom
        user_input = st.chat_input(t("type_message"))

    if user_input:
        # Check if any models are selected
        if not st.session_state.current_models:
            st.error(t("please_select_model"))
            st.stop()

        # Determine if search is needed for this message
        search_performed = False
        search_results = None

        # Check if web search is enabled
        web_search_mode = st.session_state.get("web_search_enabled", "off")
        if web_search_mode == "on":
            # Always search when web search is on
            try:
                search_manager = get_search_manager()
                with st.spinner(t("searching")):
                    import asyncio
                    search_results = asyncio.run(search_manager.perform_search(user_input))
                    search_performed = True
            except Exception as e:
                st.error(f"{t('search_error')}: {str(e)}")
                search_performed = False

        # Prepare messages for API call - WILL BE DONE INSIDE THE LOOP PER MODEL

        # Add current user message
        if st.session_state.uploaded_files:
            user_message_content = format_message_with_files(
                user_input, st.session_state.uploaded_files
            )
        else:
            user_message_content = user_input

        # Add user message to chat history with target models
        user_message = {
            "role": "user",
            "content": user_input,
            "timestamp": datetime.now().isoformat(),
            "files": st.session_state.uploaded_files.copy()
            if st.session_state.uploaded_files
            else None,
            "target_models": st.session_state.current_models.copy(), # Store which models this message is for
            "search_results": search_results.to_dict() if search_results else None
        }
        st.session_state.chat_history.append(user_message)

        # Clear uploaded files after sending
        st.session_state.uploaded_files = []

        # Display user message
        with st.chat_message("user"):
            st.write(user_input)

            # Display search results if available
            if search_results:
                with st.expander(f"🔍 {t('search_results')}", expanded=False):
                    if search_results.results:
                        for j, result in enumerate(search_results.results, 1):
                            st.markdown(f"**{j}. {result.title}**")
                            st.markdown(f"🔗 [{result.url}]({result.url})")
                            st.markdown(f"📝 {result.snippet[:300]}{'...' if len(result.snippet) > 300 else ''}")
                            if result.date:
                                st.caption(f"📅 {result.date} | 🌐 {result.source}")
                            st.divider()
                    else:
                        st.info(t("no_search_results"))

            if user_message.get("files"):
                for file_data in user_message["files"]:
                    if file_data["processed_type"] == "image":
                        st.image(file_data["content"], width=300, caption=f"Image: {file_data['name']}")
                    elif file_data["processed_type"] == "document_markdown":
                        with st.expander(f"📄 {file_data['name']} (Document Content)"):
                            st.markdown(file_data["content"])

        # Get responses from all selected models in parallel
        if st.session_state.current_models:
            # Create columns for horizontal display
            num_models = len(st.session_state.current_models)
            cols = st.columns(num_models)

            # Create queue for inter-thread communication
            response_queue = queue.Queue()

            # Create placeholders for each model
            model_placeholders = {}
            model_containers = {}
            
            # Get models to process
            thread_models = st.session_state.current_models

            # Initialize columns and placeholders
            for i, model_name in enumerate(st.session_state.current_models):
                with cols[i]:
                    st.markdown(f"### {model_name}")
                    container = st.container()
                    
                    placeholder = container.empty()
                    placeholder.markdown("⏳ Thinking...")
                    model_placeholders[model_name] = placeholder
                    model_containers[model_name] = container

            # Start threads for thread-based models simultaneously
            threads = []
            for model_name in thread_models:
                model_config = available_models[model_name]

                # Construct model-specific message history
                model_messages = [{"role": "system", "content": st.session_state.system_prompt}]

                # Add chat history filtered for this model
                for msg in st.session_state.chat_history:
                    if msg["role"] == "user":
                        # Check if this message was intended for this model
                        # Backward compatibility: if "target_models" is missing, assume it's for everyone (legacy messages)
                        if "target_models" not in msg or model_name in msg["target_models"]:
                            user_content = msg["content"]

                            # Add search results to context if available and this is the current message
                            if msg.get("search_results"):
                                search_manager = get_search_manager()
                                # Reconstruct SearchResponse from saved data
                                search_data = msg["search_results"]
                                search_response = SearchResponse(
                                    query=search_data["query"],
                                    results=[SearchResult(**result) for result in search_data["results"]],
                                    summary=search_data.get("summary")
                                )
                                search_context = search_manager.format_search_for_context(search_response)
                                user_content = f"{user_content}\n\n{search_context}"

                            if msg.get("files"):
                                model_messages.append(
                                    {
                                        "role": "user",
                                        "content": format_message_with_files(
                                            user_content, msg["files"]
                                        ),
                                    }
                                )
                            else:
                                model_messages.append({"role": "user", "content": user_content})
                    else:
                        # Assistant message - check if it belongs to this model
                        content = msg["content"]
                        # Check prefix
                        prefix = f"**{model_name}:** "
                        error_prefix = f"**{model_name} Error:** "

                        if content.startswith(prefix) or content.startswith(error_prefix):
                            # Strip model prefix
                            clean_content = re.sub(r"^\*\*.*?\*\* ", "", content)
                            model_messages.append({"role": "assistant", "content": clean_content})


                
                thread = threading.Thread(
                    target=stream_model_response,
                    args=(
                        model_config,
                        model_messages,
                        st.session_state.temperature,
                        response_queue,
                        model_name,
                    ),
                )
                thread.start()
                threads.append(thread)

            # Track responses and completion
            model_responses = {}
            completed_models = set()
            total_thread_models = len(thread_models)

            # Process responses from thread-based models
            while len(completed_models) < total_thread_models:
                try:
                    # Get response from queue with timeout
                    msg_type, model_name, *data = response_queue.get(timeout=0.1)

                    if msg_type == "chunk":
                        full_response, chunk = data
                        model_responses[model_name] = full_response
                        # Update the placeholder with streaming response
                        model_placeholders[model_name].markdown(full_response + "▌")

                    elif msg_type == "image":
                        image = data[0]
                        model_responses[model_name] = image # Store image object instead of text
                        
                        # Display image
                        model_placeholders[model_name].empty()
                        model_containers[model_name].image(image, caption=f"Generated by {model_name}", width=512)
                        
                        # Create download button
                        buf = io.BytesIO()
                        image.save(buf, format="PNG")
                        byte_im = buf.getvalue()
                        
                        model_containers[model_name].download_button(
                            label="Download Image",
                            data=byte_im,
                            file_name=f"generated_image_{int(time.time())}.png",
                            mime="image/png",
                            key=f"download_{model_name}_{int(time.time())}"
                        )

                    elif msg_type == "done":
                        full_response = data[0]
                        # Only update text if it's not an image response (image response stored as object)
                        if not isinstance(model_responses.get(model_name), Image.Image):
                             model_responses[model_name] = full_response
                             # Remove the cursor
                             model_placeholders[model_name].markdown(full_response)
                        
                        completed_models.add(model_name)

                    elif msg_type == "error":
                        error_msg = data[0]
                        model_responses[model_name] = f"Error: {error_msg}"
                        model_placeholders[model_name].error(f"❌ {error_msg}")
                        completed_models.add(model_name)

                except queue.Empty:
                    # Continue waiting for responses
                    continue

            # Wait for all threads to complete
            for thread in threads:
                thread.join()

            # Add all responses to chat history (only for thread-based models)
            for model_name in thread_models:
                if model_name in model_responses:
                    response_data = model_responses[model_name]
                    
                    if isinstance(response_data, Image.Image):
                        # Handle image response
                        assistant_message = {
                            "role": "assistant",
                            "content": f"**{model_name}:** Generated an image",
                            "files": [{
                                "name": f"generated_{int(time.time())}.png",
                                "type": "image/png",
                                "processed_type": "image",
                                "content": response_data
                            }],
                            "timestamp": datetime.now().isoformat(),
                        }
                    else:
                        # Handle text response
                        response_text = response_data
                        if not response_text.startswith("Error:"):
                            content = f"**{model_name}:** {response_text}"
                        else:
                            content = f"**{model_name} Error:** {response_text}"

                        assistant_message = {
                            "role": "assistant",
                            "content": content,
                            "timestamp": datetime.now().isoformat(),
                        }
                    
                    st.session_state.chat_history.append(assistant_message)



            # Save updated chat history to file
            save_chat_history_to_file()

        # Rerun to update the interface
        st.rerun()

# Footer
# st.markdown("🤖 **AI Chat Application** - Built with Streamlit | Supports multiple AI models and image inputs")
